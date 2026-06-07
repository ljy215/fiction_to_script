import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db import Base, get_db
from app.main import app
from app.storage import LocalFileStorage, get_file_storage


def build_text_pdf(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_start = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF\n".encode("ascii")
    )
    return bytes(content)


def build_epub(chapters: list[tuple[str, str]]) -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip")
        archive.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>""",
        )
        manifest_items = []
        spine_items = []
        for index, (title, body) in enumerate(chapters, start=1):
            item_id = f"chap{index}"
            filename = f"chapter{index}.xhtml"
            manifest_items.append(
                f'<item id="{item_id}" href="{filename}" media-type="application/xhtml+xml"/>'
            )
            spine_items.append(f'<itemref idref="{item_id}"/>')
            archive.writestr(
                f"OEBPS/{filename}",
                f"""<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <body>
    <h1>{title}</h1>
    <p>{body}</p>
  </body>
</html>""",
            )
        archive.writestr(
            "OEBPS/content.opf",
            f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <manifest>{''.join(manifest_items)}</manifest>
  <spine>{''.join(spine_items)}</spine>
</package>""",
        )
    return buffer.getvalue()


class ImportApiTest(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.storage_root = self.root / "files"
        db_path = self.root / "imports-test.sqlite3"
        self.engine = create_engine(
            f"sqlite:///{db_path}",
            connect_args={"check_same_thread": False},
        )
        Base.metadata.create_all(bind=self.engine)
        self.SessionTesting = sessionmaker(
            bind=self.engine,
            autoflush=False,
            autocommit=False,
        )

        def override_get_db():
            db = self.SessionTesting()
            try:
                yield db
            finally:
                db.close()

        def override_file_storage():
            return LocalFileStorage(self.storage_root)

        app.dependency_overrides[get_db] = override_get_db
        app.dependency_overrides[get_file_storage] = override_file_storage
        self.client = TestClient(app)

    def tearDown(self):
        app.dependency_overrides.clear()
        Base.metadata.drop_all(bind=self.engine)
        self.engine.dispose()
        self.temp_dir.cleanup()

    def auth_headers(self) -> dict[str, str]:
        payload = {
            "email": "imports@example.com",
            "password": "strong-password",
            "username": "imports",
        }
        self.client.post("/auth/register", json=payload)
        login_response = self.client.post(
            "/auth/login",
            json={"email": payload["email"], "password": payload["password"]},
        )
        token = login_response.json()["access_token"]
        return {"Authorization": f"Bearer {token}"}

    def create_project(self, headers: dict[str, str]) -> int:
        response = self.client.post(
            "/projects",
            headers=headers,
            json={"name": "导入测试项目", "script_type": "film"},
        )
        self.assertEqual(response.status_code, 201)
        return response.json()["id"]

    def test_import_pasted_text(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)

        response = self.client.post(
            f"/projects/{project_id}/imports/text",
            headers=headers,
            json={"text": "第一章\n故事开始。"},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "pasted_text")
        self.assertEqual(payload["content_text"], "第一章\n故事开始。")
        self.assertEqual(payload["content_length"], len("第一章\n故事开始。"))
        self.assertIsNone(payload["stored_file_id"])

    def test_import_pasted_text_recognizes_three_chapters(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)
        text = "\n".join(
            [
                "第一章 初遇",
                "林晚在雨夜遇见旧友。",
                "第二章 追问",
                "旧友说出失踪线索。",
                "第三章 决定",
                "林晚决定回到故乡。",
            ]
        )

        response = self.client.post(
            f"/projects/{project_id}/imports/text",
            headers=headers,
            json={"text": text},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["chapter_count"], 3)
        self.assertEqual(payload["minimum_chapters_required"], 3)
        self.assertTrue(payload["is_generation_ready"])

        chapters_response = self.client.get(
            f"/projects/{project_id}/imports/{payload['id']}/chapters",
            headers=headers,
        )

        self.assertEqual(chapters_response.status_code, 200)
        chapters = chapters_response.json()
        self.assertEqual([chapter["title"] for chapter in chapters], ["第一章 初遇", "第二章 追问", "第三章 决定"])
        self.assertEqual(chapters[0]["content_text"], "林晚在雨夜遇见旧友。")

    def test_two_chapter_import_is_not_generation_ready(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)

        response = self.client.post(
            f"/projects/{project_id}/imports/text",
            headers=headers,
            json={
                "text": "\n".join(
                    [
                        "Chapter 1: Arrival",
                        "The stranger reaches the station.",
                        "Chapter 2: Letter",
                        "A hidden letter changes the plan.",
                    ]
                )
            },
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["chapter_count"], 2)
        self.assertFalse(payload["is_generation_ready"])

    def test_import_txt_file(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)

        response = self.client.post(
            f"/projects/{project_id}/imports/txt",
            headers=headers,
            files={"file": ("novel.txt", "第一章\n文本导入。".encode("utf-8"), "text/plain")},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "txt_file")
        self.assertEqual(payload["original_filename"], "novel.txt")
        self.assertEqual(payload["content_text"], "第一章\n文本导入。")
        self.assertIsNotNone(payload["stored_file_id"])
        self.assertTrue(any(self.storage_root.rglob("*.txt")))

    def test_import_docx_file_extracts_plain_text(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)
        document = Document()
        document.add_paragraph("第一章")
        document.add_paragraph("这是 docx 正文。")
        buffer = BytesIO()
        document.save(buffer)

        response = self.client.post(
            f"/projects/{project_id}/imports/docx",
            headers=headers,
            files={
                "file": (
                    "novel.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "docx_file")
        self.assertEqual(payload["original_filename"], "novel.docx")
        self.assertEqual(payload["content_text"], "第一章\n这是 docx 正文。")
        self.assertIsNotNone(payload["stored_file_id"])

    def test_empty_docx_file_is_rejected(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)
        document = Document()
        buffer = BytesIO()
        document.save(buffer)

        response = self.client.post(
            f"/projects/{project_id}/imports/docx",
            headers=headers,
            files={
                "file": (
                    "empty.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 400)

    def test_import_pdf_file_extracts_text(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)
        pdf_bytes = build_text_pdf("Chapter 1 Arrival")

        response = self.client.post(
            f"/projects/{project_id}/imports/pdf",
            headers=headers,
            files={"file": ("novel.pdf", pdf_bytes, "application/pdf")},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "pdf_file")
        self.assertEqual(payload["original_filename"], "novel.pdf")
        self.assertIn("Chapter 1 Arrival", payload["content_text"])

    def test_empty_pdf_file_is_rejected(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)

        response = self.client.post(
            f"/projects/{project_id}/imports/pdf",
            headers=headers,
            files={"file": ("blank.pdf", b"%PDF-1.4\n%%EOF\n", "application/pdf")},
        )

        self.assertEqual(response.status_code, 400)

    def test_import_epub_file_extracts_ordered_chapters(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)
        epub_bytes = build_epub(
            [
                ("第一章 初遇", "林晚在雨夜遇见旧友。"),
                ("第二章 追问", "旧友说出失踪线索。"),
                ("第三章 决定", "林晚决定回到故乡。"),
            ]
        )

        response = self.client.post(
            f"/projects/{project_id}/imports/epub",
            headers=headers,
            files={"file": ("novel.epub", epub_bytes, "application/epub+zip")},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["source_type"], "epub_file")
        self.assertEqual(payload["chapter_count"], 3)
        self.assertTrue(payload["is_generation_ready"])

        chapters_response = self.client.get(
            f"/projects/{project_id}/imports/{payload['id']}/chapters",
            headers=headers,
        )
        self.assertEqual(chapters_response.status_code, 200)
        self.assertEqual(
            [chapter["title"] for chapter in chapters_response.json()],
            ["第一章 初遇", "第二章 追问", "第三章 决定"],
        )

    def test_invalid_epub_file_is_rejected(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)

        response = self.client.post(
            f"/projects/{project_id}/imports/epub",
            headers=headers,
            files={"file": ("broken.epub", b"not a zip", "application/epub+zip")},
        )

        self.assertEqual(response.status_code, 400)

    def test_unsupported_file_type_is_rejected(self):
        headers = self.auth_headers()
        project_id = self.create_project(headers)

        response = self.client.post(
            f"/projects/{project_id}/imports/txt",
            headers=headers,
            files={"file": ("novel.pdf", b"%PDF", "application/pdf")},
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(list(self.storage_root.rglob("*")))


if __name__ == "__main__":
    unittest.main()
