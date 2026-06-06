from pathlib import Path
from typing import Annotated

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy.orm import Session

from app.api.projects import get_owned_project
from app.auth.dependencies import get_current_user
from app.db import get_db
from app.models import SourceDocument, StoredFile, User
from app.schemas import SourceDocumentRead, TextImportCreate
from app.storage import LocalFileStorage, get_file_storage

router = APIRouter(prefix="/projects/{project_id}/imports", tags=["imports"])


def _clean_text(text: str) -> str:
    cleaned = text.replace("\ufeff", "").strip()
    if not cleaned:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Imported text cannot be empty",
        )
    return cleaned


def _decode_text_file(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return _clean_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="TXT file must use UTF-8 or GB18030 compatible text encoding",
    )


def _require_extension(filename: str, extension: str) -> str:
    safe_name = Path(filename or "").name
    if Path(safe_name).suffix.lower() != extension:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only {extension} files are supported in this import step",
        )
    return safe_name


def _save_uploaded_file(
    db: Session,
    storage: LocalFileStorage,
    file: UploadFile,
    owner_id: int,
    project_id: int,
    filename: str,
) -> tuple[StoredFile, Path]:
    stored_payload = storage.save_upload(
        stream=file.file,
        original_filename=filename,
        content_type=file.content_type,
        owner_id=owner_id,
    )
    stored_file = StoredFile(
        owner_id=owner_id,
        project_id=project_id,
        **stored_payload.__dict__,
    )
    db.add(stored_file)
    db.flush()
    return stored_file, storage.absolute_path(stored_payload.relative_path)


def _extract_docx_text(path: Path) -> str:
    try:
        document = Document(path)
    except PackageNotFoundError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="DOCX file cannot be parsed",
        ) from exc

    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return _clean_text("\n".join(parts))


def _create_source_document(
    db: Session,
    owner_id: int,
    project_id: int,
    source_type: str,
    text: str,
    original_filename: str | None = None,
    stored_file_id: int | None = None,
) -> SourceDocument:
    document = SourceDocument(
        owner_id=owner_id,
        project_id=project_id,
        stored_file_id=stored_file_id,
        source_type=source_type,
        original_filename=original_filename,
        content_text=text,
        content_length=len(text),
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.post("/text", response_model=SourceDocumentRead, status_code=status.HTTP_201_CREATED)
def import_pasted_text(
    project_id: int,
    payload: TextImportCreate,
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
):
    get_owned_project(db, project_id, current_user.id)
    text = _clean_text(payload.text)
    return _create_source_document(
        db=db,
        owner_id=current_user.id,
        project_id=project_id,
        source_type="pasted_text",
        text=text,
    )


@router.post("/txt", response_model=SourceDocumentRead, status_code=status.HTTP_201_CREATED)
def import_txt_file(
    project_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
    storage: Annotated[LocalFileStorage, Depends(get_file_storage)],
    file: UploadFile = File(...),
):
    get_owned_project(db, project_id, current_user.id)

    filename = _require_extension(file.filename or "", ".txt")
    stored_file, stored_path = _save_uploaded_file(
        db=db,
        storage=storage,
        file=file,
        owner_id=current_user.id,
        project_id=project_id,
        filename=filename,
    )

    text = _decode_text_file(stored_path.read_bytes())
    document = SourceDocument(
        owner_id=current_user.id,
        project_id=project_id,
        stored_file_id=stored_file.id,
        source_type="txt_file",
        original_filename=stored_file.original_filename,
        content_text=text,
        content_length=len(text),
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.post("/docx", response_model=SourceDocumentRead, status_code=status.HTTP_201_CREATED)
def import_docx_file(
    project_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
    storage: Annotated[LocalFileStorage, Depends(get_file_storage)],
    file: UploadFile = File(...),
):
    get_owned_project(db, project_id, current_user.id)

    filename = _require_extension(file.filename or "", ".docx")
    stored_file, stored_path = _save_uploaded_file(
        db=db,
        storage=storage,
        file=file,
        owner_id=current_user.id,
        project_id=project_id,
        filename=filename,
    )

    text = _extract_docx_text(stored_path)
    document = SourceDocument(
        owner_id=current_user.id,
        project_id=project_id,
        stored_file_id=stored_file.id,
        source_type="docx_file",
        original_filename=stored_file.original_filename,
        content_text=text,
        content_length=len(text),
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document
